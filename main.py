from zipfile import ZipFile
from fastapi import FastAPI, HTTPException, Response
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Dict, Any
import logging
import os
import tempfile
from docx import Document
import shutil
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import subprocess
from fastapi.responses import FileResponse
from datetime import datetime, timedelta
import hijri_converter  # pip install hijri-converter

app = FastAPI()

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # In production, replace with specific origins
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Define data models
class Person(BaseModel):
    name: str
    nationality: str
    id_number: str

class GregorianDate(BaseModel):
    date: str  # Format: "YYYY-MM-DD" or ISO format date string

class GetPassData(BaseModel):
    people: List[Person]
    dates: List[GregorianDate]


def embed_fonts(doc):
    """
    Attempt to add font embedding settings to the Word document
    """
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        # Create new settings element if it doesn't exist
        if not doc.settings._element:
            doc._part.add_settings_part()
        
        settings = doc.settings._element
        
        # Try to add the embedSystemFonts setting
        embed_fonts_elm = OxmlElement('w:embedSystemFonts')
        embed_fonts_elm.set(qn('w:val'), 'true')
        settings.append(embed_fonts_elm)
        
        # Try to add embedTrueTypeFonts setting
        embed_tt_fonts = OxmlElement('w:embedTrueTypeFonts')
        embed_tt_fonts.set(qn('w:val'), 'true')
        settings.append(embed_tt_fonts)
        
        # Add saveSubsetFonts setting
        save_subset = OxmlElement('w:saveSubsetFonts')
        save_subset.set(qn('w:val'), 'true')
        settings.append(save_subset)
        
        logger.info("Added font embedding settings to document")
        return True
    except Exception as e:
        logger.error(f"Failed to set font embedding: {str(e)}")
        return False

# Arabic day names
ARABIC_DAY_NAMES = ["الأحد", "الإثنين", "الثلاثاء", "الأربعاء", "الخميس", "الجمعة", "السبت"]

def convert_to_hijri(date_str):
    """
    Convert a Gregorian date string to Hijri date
    Returns a tuple (hijri_date_str, arabic_day_name, gregorian_date_str)
    """
    try:
        # Parse the date string to a datetime object
        date_obj = datetime.fromisoformat(date_str.replace('Z', '+00:00'))
        
        # Get Arabic day name
        day_of_week = date_obj.weekday()
        # Convert to Arabic convention (Sunday=0, Monday=1, etc.)
        arabic_day_index = (day_of_week + 1) % 7
        arabic_day_name = ARABIC_DAY_NAMES[arabic_day_index]
        
        # Convert to Hijri using hijri-converter
        hijri = hijri_converter.convert.Gregorian(
            date_obj.year, 
            date_obj.month, 
            date_obj.day
        ).to_hijri()
        
        # Format dates as strings
        hijri_date_str = f"{hijri.day:02d}/{hijri.month:02d}/{hijri.year}"
        gregorian_date_str = f"{date_obj.day:02d}/{date_obj.month:02d}/{date_obj.year}"
        
        # Log the conversion for debugging
        logger.info(f"Converted date: Gregorian {gregorian_date_str} to Hijri {hijri_date_str}, Day: {arabic_day_name}")
        
        return hijri_date_str, arabic_day_name, gregorian_date_str
    
    except Exception as e:
        logger.error(f"Error converting date: {e}")
        raise ValueError(f"Invalid date format: {date_str}")

def set_cell_style(cell):
    # تعيين الخط Arial (Body CS) والحجم 15
    cell.text = cell.text.strip()  # إزالة أي فراغات زائدة
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Arial (Body CS)"  # تعيين الخط
            run.font.size = Pt(15)  # تعيين الحجم
            run.bold = True
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # توسيط النص داخل الخلية

def process_document(date_info, people, template_file, output_file):
    """
    Process the document with the given data
    date_info is a tuple (hijri_date_str, arabic_day_name, gregorian_date_str)
    """
    # نسخ الملف إلى ملف جديد
    shutil.copy(template_file, output_file)
    doc = Document(output_file)

    # استخدم التاريخ والبيانات المرسلة
    hijri_date_str, day, gregorian_date_str = date_info

    # Prepare replacements based on the date data
    replacements = {
        "(اليوم)": day,
        "[D]": hijri_date_str.split('/')[0],
        "[M]": hijri_date_str.split('/')[1],
        "[Y]": hijri_date_str.split('/')[2],
        "[d]": gregorian_date_str.split('/')[0],
        "[m]": gregorian_date_str.split('/')[1],
        "[yyyy]": gregorian_date_str.split('/')[2],
    }

    num_people = len(people)
    # استبدال العدد في المكان المخصص للأرقام (ع)
    replacements["الموضح هوياتهم بالبيان المرفق وعددهم (ع)"] = f"الموضح هوياتهم بالبيان المرفق وعددهم ({num_people:02d}) ☒"

    if num_people == 1:
        # حالة الزائر الواحد: استبدال بيانات الزائر الأول وتفريغ بيانات الزائر الثاني
        person = people[0]
        replacements["(الزائر1)"] = person.name
        replacements["(الهويه1)"] = person.id_number
        replacements["(الجنسيه1)"] = person.nationality
        replacements["(الزائر2)"] = ""
        replacements["(الهويه2)"] = ""
        replacements["(الجنسيه2)"] = ""
        replacements["(اولهم)"] = ""
        replacements["(اخرهم)"] = ""
    elif num_people == 2:
        # حالة وجود زائرين
        for i, person in enumerate(people, start=1):
            replacements[f"(الزائر{i})"] = person.name
            replacements[f"(الهويه{i})"] = person.id_number
            replacements[f"(الجنسيه{i})"] = person.nationality
        replacements["(اولهم)"] = ""
        replacements["(اخرهم)"] = ""
    else:
        # حالة أكثر من زائرين - تحديث: ترك حقول الزائرين الأول والثاني فارغة
        replacements["(الزائر1)"] = ""
        replacements["(الهويه1)"] = ""
        replacements["(الجنسيه1)"] = ""
        replacements["(الزائر2)"] = ""
        replacements["(الهويه2)"] = ""
        replacements["(الجنسيه2)"] = ""
        # إضافة أولهم وآخرهم
        replacements["(اولهم)"] = people[0].name
        replacements["(اخرهم)"] = people[-1].name

    # استبدال النصوص داخل الجداول وتعديل التنسيقات
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, str(value))
                
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        text = run.text.strip()
                        if "لموضح هوياتهم بالبيان" in text:
                            run.font.size = Pt(15)
                            run.font.bold = False
                            run.font.name = 'Times New Roman (Headings CS)'
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                        elif text in {replacements.get("[D]"), replacements.get("[M]"), 
                                      replacements.get("[Y]"), replacements.get("[d]"), 
                                      replacements.get("[m]"), replacements.get("[yyyy]")}:
                            run.font.size = Pt(8)
                            run.font.bold = True  
                            run.font.name = 'Arial (Body CS)'
                            if text in {replacements.get("[d]"), replacements.get("[D]")}:
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            elif text in {replacements.get("[M]"), replacements.get("[m]")}:
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER       
                            elif text in {replacements.get("[Y]"), replacements.get("[yyyy]")}:
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                        elif text in {p.name for p in people} | {p.nationality for p in people} | {p.id_number for p in people} | {day}:
                            run.font.size = Pt(15)
                            run.font.bold = True  
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT 
                        else:
                            run.font.size = Pt(15)
                            run.font.bold = True
                            run.font.name = 'Times New Roman (Headings CS)'

    # إذا كان عدد الزوار أكثر من 2، أضف جميع الزوار إلى الجدول الثاني
    if num_people > 2:
        # الوصول إلى الجدول الثاني في المستند
        table = doc.tables[1]  # الجدول الثاني في المستند
        
        # إضافة البيانات إلى الجدول الثاني لجميع الزوار
        for i, person in enumerate(people, start=1):
            # تأكد من أن هناك صفوف كافية
            if i <= len(table.rows) - 1:  # نبدأ من الصف الأول في الجدول الثاني
                row = table.rows[i]
            else:
                # إضافة صف جديد إذا لم يكن هناك صفوف كافية
                row = table.add_row()
            
            # ملء البيانات في الصف
            row.cells[0].text = person.id_number
            row.cells[1].text = person.nationality
            row.cells[2].text = person.name
            
            # تطبيق التنسيق على الخلايا
            for j in range(3):
                set_cell_style(row.cells[j])


    embed_fonts(doc)

    doc.save(output_file)
    logger.info(f"تم إنشاء الملف بنجاح: {output_file}")
    return output_file

def convert_to_pdf(docx_file, pdf_file):
    """
    Convert a Word document to PDF using unoconv
    Returns True if successful, False otherwise
    """
    try:
        # Use unoconv for conversion (must be installed on the server)
        cmd = ['unoconv', '-f', 'pdf', '-o', pdf_file, docx_file]
        
        process = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        stdout, stderr = process.communicate()
        
        if process.returncode != 0:
            logger.error(f"Error converting to PDF: {stderr.decode()}")
            return False
        
        # Check if the output file exists
        if not os.path.exists(pdf_file):
            logger.error(f"PDF file was not created: {pdf_file}")
            return False
            
        return True
    except Exception as e:
        logger.error(f"Error in PDF conversion: {str(e)}")
        return False

def merge_pdfs(pdf_files, output_file):
    """
    Merge multiple PDF files into one using PyPDF2
    Returns True if successful, False otherwise
    """
    try:
        # Use PDFMerger from PyPDF2
        import PyPDF2
        
        merger = PyPDF2.PdfMerger()
        
        for pdf in pdf_files:
            merger.append(pdf)
            
        merger.write(output_file)
        merger.close()
        
        return True
    except Exception as e:
        logger.error(f"Error merging PDFs: {str(e)}")
        return False

def create_docx_zip(docx_files, output_zip):
    """
    Create a ZIP file containing multiple DOCX files
    """
    try:
        with ZipFile(output_zip, 'w') as zipf:
            for docx_file in docx_files:
                zipf.write(docx_file, os.path.basename(docx_file))
        return True
    except Exception as e:
        logger.error(f"Error creating ZIP file: {str(e)}")
        return False

def merge_docx_files(docx_files, output_file):
    """
    Merge multiple Word documents into one
    """
    try:
        # Use the first document as the base
        merged_doc = Document(docx_files[0])
        
        # Add a page break after the first document
        merged_doc.add_page_break()
        
        # Append each additional document
        for i in range(1, len(docx_files)):
            doc = Document(docx_files[i])
            
            # Copy all elements from the document
            for element in doc.element.body:
                merged_doc.element.body.append(element)
            
            # Add a page break after each document (except the last one)
            if i < len(docx_files) - 1:
                merged_doc.add_page_break()
        
        # Save the merged document
        merged_doc.save(output_file)
        return True
    except Exception as e:
        logger.error(f"Error merging Word documents: {str(e)}")
        return False

# @app.post("/generate-getpass/")
# async def generate_getpass(data: GetPassData):
#     try:
#         # Create a temporary directory for working files
#         with tempfile.TemporaryDirectory() as temp_dir:
#             template_file = "GETPASS.docx"  # Path to your template file
#             docx_files = []
#             pdf_files = []
#             conversion_success = True
            
#             # Process each date entry
#             for i, date_entry in enumerate(data.dates):
#                 try:
#                     # Fix timezone issues by ensuring we have a complete datetime
#                     date_str = date_entry.date
#                     if 'T' not in date_str:
#                         # If only a date is provided (YYYY-MM-DD), add time component
#                         date_str = f"{date_str}T12:00:00"
                    
#                     # Convert Gregorian date to Hijri with correct handling
#                     date_info = convert_to_hijri(date_str)
                    
#                     docx_file = os.path.join(temp_dir, f"getpass_{i}.docx")
#                     pdf_file = os.path.join(temp_dir, f"getpass_{i}.pdf")
                    
#                     # Process the document with converted date
#                     processed_docx = process_document(date_info, data.people, template_file, docx_file)
#                     docx_files.append(processed_docx)
                    
#                     # Try to convert to PDF
#                     if convert_to_pdf(docx_file, pdf_file):
#                         pdf_files.append(pdf_file)
#                     else:
#                         conversion_success = False
#                         logger.warning(f"Failed to convert {docx_file} to PDF")
                        
#                 except ValueError as e:
#                     logger.error(f"Invalid date format: {e}")
#                     raise HTTPException(status_code=400, detail=f"Invalid date format: {str(e)}")
            
#             # Create output directory
#             output_dir = "output"
#             os.makedirs(output_dir, exist_ok=True)
            
#             # If PDF conversion was successful and there are multiple PDFs
#             if conversion_success and len(pdf_files) > 1:
#                 final_pdf_path = os.path.join(output_dir, "getpass.pdf")
#                 if merge_pdfs(pdf_files, final_pdf_path):
#                     return FileResponse(
#                         path=final_pdf_path,
#                         filename="getpass.pdf",
#                         media_type="application/pdf"
#                     )
#                 else:
#                     # If merging PDFs fails, fall back to Word documents
#                     conversion_success = False
            
#             # If PDF conversion was successful but there's only one PDF
#             elif conversion_success and len(pdf_files) == 1:
#                 final_pdf_path = os.path.join(output_dir, "getpass.pdf")
#                 shutil.copy2(pdf_files[0], final_pdf_path)
#                 return FileResponse(
#                     path=final_pdf_path,
#                     filename="getpass.pdf",
#                     media_type="application/pdf"
#                 )
            
#             # If PDF conversion failed or PDF merging failed, return merged Word document
#             if not conversion_success:
#                 if len(docx_files) == 1:
#                     # Return a single Word document
#                     final_docx_path = os.path.join(output_dir, "getpass.docx")
#                     shutil.copy2(docx_files[0], final_docx_path)
#                     return FileResponse(
#                         path=final_docx_path,
#                         filename="getpass.docx",
#                         media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#                     )
#                 else:
#                     # Merge multiple Word documents instead of creating a ZIP
#                     final_docx_path = os.path.join(output_dir, "getpass.docx")
#                     if merge_docx_files(docx_files, final_docx_path):
#                         return FileResponse(
#                             path=final_docx_path,
#                             filename="getpass.docx",
#                             media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#                         )
#                     else:
#                         # If merging Word documents fails, fall back to ZIP
#                         zip_path = os.path.join(output_dir, "getpass_documents.zip")
#                         with ZipFile(zip_path, 'w') as zipf:
#                             for docx_file in docx_files:
#                                 zipf.write(docx_file, os.path.basename(docx_file))
#                         return FileResponse(
#                             path=zip_path,
#                             filename="getpass_documents.zip",
#                             media_type="application/zip"
#                         )
            
#             # If we got here, something unexpected happened
#             raise HTTPException(status_code=500, detail="Failed to generate documents")
            
#     except Exception as e:
#         logger.error(f"Error processing request: {str(e)}")
#         raise HTTPException(status_code=500, detail=f"Internal Server Error: {str(e)}")


@app.post("/generate-getpass/")
async def generate_getpass(data: GetPassData):
    try:
        # Create a temporary directory for working files
        with tempfile.TemporaryDirectory() as temp_dir:
            template_file = "GETPASS.docx"  # Path to your template file
            docx_files = []
            pdf_files = []
            conversion_success = True
            
            # Convert each date and process the document
            for i, date_entry in enumerate(data.dates):
                try:
                    # Convert Gregorian date to Hijri
                    date_info = convert_to_hijri(date_entry.date)
                    
                    docx_file = os.path.join(temp_dir, f"getpass_{i}.docx")
                    pdf_file = os.path.join(temp_dir, f"getpass_{i}.pdf")
                    
                    # Process the document with converted date
                    processed_docx = process_document(date_info, data.people, template_file, docx_file)
                    docx_files.append(processed_docx)
                    
                    # Try to convert to PDF
                    if convert_to_pdf(docx_file, pdf_file):
                        pdf_files.append(pdf_file)
                    else:
                        conversion_success = False
                        logger.warning(f"Failed to convert {docx_file} to PDF")
                        
                except ValueError as e:
                    logger.error(f"Invalid date format: {e}")
                    raise HTTPException(status_code=400, detail=f"Invalid date format: {str(e)}")
            
            # Create output directory
            output_dir = "output"
            os.makedirs(output_dir, exist_ok=True)
            
            # If PDF conversion was successful and there are multiple PDFs
            if conversion_success and len(pdf_files) > 1:
                final_pdf_path = os.path.join(output_dir, "getpass.pdf")
                if merge_pdfs(pdf_files, final_pdf_path):
                    return FileResponse(
                        path=final_pdf_path,
                        filename="getpass.pdf",
                        media_type="application/pdf"
                    )
                else:
                    # If merging PDFs fails, fall back to Word documents
                    conversion_success = False
            
            # If PDF conversion was successful but there's only one PDF
            elif conversion_success and len(pdf_files) == 1:
                final_pdf_path = os.path.join(output_dir, "getpass.pdf")
                shutil.copy2(pdf_files[0], final_pdf_path)
                return FileResponse(
                    path=final_pdf_path,
                    filename="getpass.pdf",
                    media_type="application/pdf"
                )
            
            # If PDF conversion failed or PDF merging failed, return Word documents
            if not conversion_success:
                if len(docx_files) == 1:
                    # Return a single Word document
                    final_docx_path = os.path.join(output_dir, "getpass.docx")
                    shutil.copy2(docx_files[0], final_docx_path)
                    return FileResponse(
                        path=final_docx_path,
                        filename="getpass.docx",
                        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                else:
                    # Create a ZIP file with all Word documents
                    zip_path = os.path.join(output_dir, "getpass_documents.zip")
                    if create_docx_zip(docx_files, zip_path):
                        return FileResponse(
                            path=zip_path,
                            filename="getpass_documents.zip",
                            media_type="application/zip"
                        )
            
            # If we got here, something unexpected happened
            raise HTTPException(status_code=500, detail="Failed to generate documents")
            
    except Exception as e:
        logger.error(f"Error processing request: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Internal Server Error: {str(e)}")


@app.get("/")
async def root():
    return {"message": "GetPass API is running. Use /generate-getpass/ endpoint to generate getpass documents."}